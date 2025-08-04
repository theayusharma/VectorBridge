import asyncio
import email
import hashlib
import io
import json
import logging
import os
import re
from concurrent.futures import ThreadPoolExecutor
from contextlib import asynccontextmanager
from functools import partial
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from urllib.parse import urlparse

import docx
import httpx
import pdfplumber
import textract
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, Security, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.httpsredirect import HTTPSRedirectMiddleware
from fastapi.responses import JSONResponse
from fastapi.security import APIKeyHeader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS
from pydantic import BaseModel, Field, HttpUrl

from geminiloadbalance import LLMGenerationError, generate_text_with_retry

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

load_dotenv()


class Config:
    MAX_FILE_SIZE = 30 * 1024 * 1024  # 30MB
    REQUEST_TIMEOUT = 30.0
    CHUNK_SIZE = 800
    CHUNK_OVERLAP = 150
    TOP_K_CHUNKS = 5
    EMBEDDING_MODEL = "all-MiniLM-L6-v2"
    GEMINI_MODEL = "gemini-1.5-flash"
    MAX_LLM_RETRIES = 3

    # context parameters
    CONTEXT_EXPANSION_RATIO = 1.5  # expand context by 50%
    SEMANTIC_SIMILARITY_THRESHOLD = 0.3
    KEYWORD_BOOST_FACTOR = 1.2

    # cache
    EMBEDDING_MODEL_PATH = f"./models/{EMBEDDING_MODEL}"
    CACHE_DIR = Path("cache")
    DOCUMENT_CACHE_DIR = CACHE_DIR / "document"
    VECTORDB_CACHE_DIR = CACHE_DIR / "vectordb"
    METADATA_CACHE_DIR = CACHE_DIR / "metadata"


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("Loading embedding model...")
    app.state.embeddings = SentenceTransformerEmbeddings(
        model_name=Config.EMBEDDING_MODEL, cache_folder=Config.EMBEDDING_MODEL_PATH
    )
    # Create cache directories if they don't exist
    Config.DOCUMENT_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    Config.VECTORDB_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    Config.METADATA_CACHE_DIR.mkdir(parents=True, exist_ok=True)

    yield
    logger.info("Shutting down...")


# FastAPI app
api_key_header = APIKeyHeader(name="Authorization", auto_error=False)
app = FastAPI(title="HackRx RAG API", version="0.3.0", lifespan=lifespan)

if os.getenv("ENVIRONMENT") == "production":
    from fastapi.middleware.httpsredirect import HTTPSRedirectMiddleware

    app.add_middleware(HTTPSRedirectMiddleware)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# data structures
class ChunkMetadata:
    def __init__(
        self,
        chunk_id: str,
        page_num: Optional[int] = None,
        section: Optional[str] = None,
        keywords: Optional[List[str]] = None,
        chunk_type: str = "text",
        importance_score: float = 1.0,
    ):
        self.chunk_id = chunk_id
        self.page_num = page_num
        self.section = section
        self.keywords = keywords or []
        self.chunk_type = chunk_type
        self.importance_score = importance_score


class DocumentStructure:
    def __init__(self):
        self.sections: Dict[str, List[str]] = {}
        self.headers: List[Tuple[str, int]] = []  # (header_text, level)
        self.tables: List[Dict[str, Any]] = []
        self.key_value_pairs: Dict[str, str] = {}
        self.definitions: Dict[str, str] = {}


# custom exceptions
class DocumentProcessingError(Exception):
    pass


# exception handlers
@app.exception_handler(DocumentProcessingError)
async def document_processing_exception_handler(
    request: Request, exc: DocumentProcessingError
):
    return JSONResponse(
        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
        content={"detail": str(exc)},
    )


@app.exception_handler(LLMGenerationError)
async def llm_generation_exception_handler(request: Request, exc: LLMGenerationError):
    return JSONResponse(
        status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
        content={"detail": str(exc)},
    )


# auth check
async def get_api_key(api_key_header: str = Security(api_key_header)):
    if api_key_header is None:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Authorization header is missing.",
        )

    parts = api_key_header.split()
    if len(parts) != 2 or parts[0].lower() != "bearer":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid Authorization header format. Must be 'Bearer <key>'.",
        )

    token = parts[1]
    if token != os.getenv("HACKRX_AUTH_TOKEN"):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid Bearer token",
        )
    return token


# Data models
class HackRxRequest(BaseModel):
    documents: HttpUrl
    questions: List[str]

    chunk_size: Optional[int] = Field(
        default=Config.CHUNK_SIZE,
        ge=100,
        le=2000,
        description="Size of text chunks for processing",
    )
    chunk_overlap: Optional[int] = Field(
        default=Config.CHUNK_OVERLAP,
        ge=0,
        le=500,
        description="Overlap between text chunks",
    )
    top_k: Optional[int] = Field(
        default=Config.TOP_K_CHUNKS,
        ge=1,
        le=15,
        description="Number of relevant chunks to retrieve",
    )


class HackRxResponse(BaseModel):
    answers: List[str]
    processing_time: float
    document_metadata: Optional[Dict[str, Any]]
    context_metadata: Optional[Dict[str, Any]]


# keyword extraction
def extract_keywords(text: str) -> List[str]:
    # insurance-specific keywords
    insurance_keywords = [
        "premium",
        "deductible",
        "coverage",
        "exclusion",
        "claim",
        "policy",
        "benefit",
        "liability",
        "copay",
        "coinsurance",
        "network",
        "provider",
        "pre-existing",
        "waiting period",
        "grace period",
        "effective date",
        "termination",
        "renewal",
        "rider",
        "endorsement",
    ]

    # capitalized terms
    capitalized_terms = re.findall(r"\b[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*\b", text)

    # numbers with currency or percentages
    financial_terms = re.findall(r"[\$₹]\s?[\d,]+(?:\.\d{2})?|\d+(?:\.\d+)?%", text)

    # dates
    date_patterns = re.findall(
        r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\w+\s+\d{1,2},?\s+\d{4}\b", text
    )

    # combine all extracted terms
    keywords = []
    keywords.extend([kw for kw in insurance_keywords if kw.lower() in text.lower()])
    keywords.extend(capitalized_terms[:5])  # limit to 5
    keywords.extend(financial_terms)
    keywords.extend(date_patterns)

    return list(set(keywords))


# PDF parser
def parse_pdf(content: io.BytesIO) -> Tuple[str, Dict[str, Any], DocumentStructure]:
    full_text = []
    metadata = {"pages": 0, "tables": 0, "sections": 0}
    doc_structure = DocumentStructure()

    try:
        with pdfplumber.open(content) as pdf:
            metadata["pages"] = len(pdf.pages)
            current_section = "Introduction"

            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text(x_tolerance=2) or ""

                # headers (lines that are likely section headers)
                lines = page_text.split("\n")
                for line in lines:
                    line = line.strip()
                    if line and (
                        line.isupper()
                        or re.match(r"^[A-Z][A-Za-z\s]+:$", line)
                        or re.match(r"^\d+\.\s+[A-Z]", line)
                    ):
                        doc_structure.headers.append((line, 1))
                        current_section = line
                        metadata["sections"] += 1

                # add to chunks
                if current_section not in doc_structure.sections:
                    doc_structure.sections[current_section] = []
                doc_structure.sections[current_section].append(page_text)

                full_text.append(
                    f"--- Page {i+1} | Section: {current_section} ---\n{page_text}"
                )

                # tables
                tables = page.extract_tables()
                if tables:
                    metadata["tables"] += len(tables)
                    for j, table in enumerate(tables):
                        table_metadata = {
                            "page": i + 1,
                            "table_id": f"table_{i+1}_{j+1}",
                            "rows": len(table),
                            "cols": len(table[0]) if table else 0,
                        }
                        doc_structure.tables.append(table_metadata)

                        # conversion to markdown
                        if table and len(table) > 1:
                            headers = table[0]
                            markdown_table = (
                                "| "
                                + " | ".join(
                                    str(cell) if cell else "" for cell in headers
                                )
                                + " |\n"
                            )
                            markdown_table += "|" + "---|" * len(headers) + "\n"

                            for row in table[1:]:
                                markdown_table += (
                                    "| "
                                    + " | ".join(
                                        str(cell) if cell else "" for cell in row
                                    )
                                    + " |\n"
                                )

                            full_text.append(
                                f"\n--- Table {j+1} on Page {i+1} ---\n{markdown_table}\n"
                            )

                # key-value pairs
                kv_matches = re.findall(r"([A-Z][A-Za-z\s]+):\s*([^\n]+)", page_text)
                for key, value in kv_matches:
                    doc_structure.key_value_pairs[key.strip()] = value.strip()
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse PDF: {str(e)}")

    return "\n".join(full_text), metadata, doc_structure


# DOCX parser
def parse_docx(content: io.BytesIO) -> Tuple[str, Dict[str, Any], DocumentStructure]:
    try:
        doc = docx.Document(content)
        doc_structure = DocumentStructure()
        full_text = []
        current_section = "Introduction"

        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
        }

        for para in doc.paragraphs:
            if para.text.strip():  # todo check errors from lsp
                # header
                if para.style.name.startswith("Heading") or (
                    para.text.isupper() and len(para.text) < 100
                ):
                    level = 1
                    if "Heading" in para.style.name:
                        level = (
                            int(para.style.name.split()[-1])
                            if para.style.name.split()[-1].isdigit()
                            else 1
                        )
                    doc_structure.headers.append((para.text, level))
                    current_section = para.text

                # add to section
                if current_section not in doc_structure.sections:
                    doc_structure.sections[current_section] = []
                doc_structure.sections[current_section].append(para.text)

                full_text.append(f"[Section: {current_section}] {para.text}")

        # tables
        for i, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))

            table_metadata = {
                "table_id": f"docx_table_{i+1}",
                "rows": len(table.rows),
                "cols": len(table.rows[0].cells) if table.rows else 0,
            }
            doc_structure.tables.append(table_metadata)
            full_text.append(f"\n--- Table {i+1} ---\n" + "\n".join(table_text) + "\n")

        return "\n".join(full_text), metadata, doc_structure
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse DOCX: {str(e)}")


def parse_doc(
    content: io.BytesIO,
) -> Tuple[str, Dict[str, Any], Optional[DocumentStructure]]:
    temp_file_path = "temp_file.doc"
    try:
        with open(temp_file_path, "wb") as f:
            f.write(content.read())

        text = textract.process(temp_file_path).decode("utf-8")
        return text, {"format": "DOC"}, None
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse DOC: {str(e)}")
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)


def parse_eml(
    content: io.BytesIO,
) -> Tuple[str, Dict[str, Any], Optional[DocumentStructure]]:
    try:
        msg = email.message_from_bytes(content.read())
        metadata = {
            "subject": msg["subject"],
            "from": msg["from"],
            "to": msg["to"],
            "date": msg["date"],
            "content_type": msg.get_content_type(),
        }

        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                if ctype == "text/plain":
                    body = part.get_payload(decode=True).decode(errors="ignore")
                    break
        else:
            body = msg.get_payload(decode=True).decode(errors="ignore")

        return (
            f"Subject: {msg['subject']}\nFrom: {msg['from']}\nTo: {msg['to']}\n\n{body}",
            metadata,
            None,
        )
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse EML: {str(e)}")


def parse_txt(
    content: io.BytesIO,
) -> Tuple[str, Dict[str, Any], Optional[DocumentStructure]]:
    try:
        text = content.read().decode("utf-8", errors="ignore")
        return text, {"size": len(text)}, None
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse TXT: {str(e)}")


# parser mapping
PARSER_MAPPING: Dict[str, callable] = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_doc,
    ".eml": parse_eml,
    ".txt": parse_txt,
}


# doc downloader
async def get_document_content_from_url(
    url: Union[str, HttpUrl],
) -> Tuple[io.BytesIO, str]:
    logger.info(f"Downloading document from {url}")

    url_str = str(url) if isinstance(url, HttpUrl) else url

    try:
        async with httpx.AsyncClient(timeout=Config.REQUEST_TIMEOUT) as client:
            head_response = await client.head(url_str, follow_redirects=True)
            head_response.raise_for_status()

            content_length = head_response.headers.get("content-length")
            if content_length and int(content_length) > Config.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"Document exceeds maximum size of {Config.MAX_FILE_SIZE/1024/1024}MB",
                )

            response = await client.get(url_str, follow_redirects=True)
            response.raise_for_status()

            file_ext = Path(urlparse(url_str).path).suffix.lower()
            if not file_ext or file_ext not in PARSER_MAPPING:
                file_ext = ".txt"

            content = io.BytesIO(await response.aread())
            if content.getbuffer().nbytes > Config.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"Document exceeds maximum size of {Config.MAX_FILE_SIZE/1024/1024}MB",
                )

            return content, file_ext
    except httpx.RequestError as e:
        logger.error(f"Failed to download document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to download document: {str(e)}",
        )


# text processing with intelligent chunking
def clean_and_chunk_text(
    raw_text: str,
    chunk_size: int = Config.CHUNK_SIZE,
    chunk_overlap: int = Config.CHUNK_OVERLAP,
) -> Tuple[List[str], List[ChunkMetadata]]:
    logger.info("Processing and chunking of text")

    # text cleaning
    cleaned_text = re.sub(r"\n{3,}", "\n\n", raw_text).strip()
    cleaned_text = re.sub(r"[ \t]{2,}", " ", cleaned_text)

    # remove excessive spacing, preserve structure
    cleaned_text = re.sub(r"([.!?])\s*\n\s*([A-Z])", r"\1\n\n\2", cleaned_text)

    # text splitter with custom separators for semantic chunking
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=[
            "\n\n--- Table",  # table boundaries
            "\n\n--- Page",  # page boundaries
            "\n\n[Section:",  # section boundaries
            "\n\n",  # paragraph breaks
            "\n",  # line breaks
            ". ",  # sentence breaks
            " ",  # word breaks
            "",  # character breaks
        ],
    )

    chunks = text_splitter.split_text(cleaned_text)

    if not chunks:
        raise DocumentProcessingError(
            "Could not extract meaningful text chunks from document"
        )

    # metadata for each chunk
    chunk_metadata = []
    for i, chunk in enumerate(chunks):
        # page number
        page_match = re.search(r"--- Page (\d+)", chunk)
        page_num = int(page_match.group(1)) if page_match else None

        # section
        section_match = re.search(r"\[Section: ([^\]]+)\]", chunk)
        section = section_match.group(1) if section_match else "General"

        # keywords
        keywords = extract_keywords(chunk)

        # determine chunk type
        chunk_type = "text"
        if "--- Table" in chunk:
            chunk_type = "table"
        elif any(
            keyword in chunk.lower() for keyword in ["exclusion", "coverage", "benefit"]
        ):
            chunk_type = "policy_rule"

        # importance score (based on keywords and structure)
        importance_score = 1.0
        if chunk_type == "policy_rule":
            importance_score = 1.5
        if len(keywords) > 3:
            importance_score *= 1.2

        metadata = ChunkMetadata(
            chunk_id=f"chunk_{i}",
            page_num=page_num,
            section=section,
            keywords=keywords,
            chunk_type=chunk_type,
            importance_score=importance_score,
        )
        chunk_metadata.append(metadata)

    logger.info(f"Split text into {len(chunks)} chunks")
    return chunks, chunk_metadata


# caching
def get_document_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def save_metadata(
    doc_hash: str,
    metadata: Dict[str, Any],
    doc_structure: Optional[DocumentStructure] = None,
    chunk_metadata: Optional[List[ChunkMetadata]] = None,
):
    cache_data = {
        "metadata": metadata,
        "doc_structure": doc_structure.__dict__ if doc_structure else None,
        "chunk_metadata": (
            [
                {
                    "chunk_id": cm.chunk_id,
                    "page_num": cm.page_num,
                    "section": cm.section,
                    "keywords": cm.keywords,
                    "chunk_type": cm.chunk_type,
                    "importance_score": cm.importance_score,
                }
                for cm in chunk_metadata
            ]
            if chunk_metadata
            else None
        ),
    }

    with open(Config.METADATA_CACHE_DIR / f"{doc_hash}.json", "w") as f:
        json.dump(cache_data, f, indent=2)


def load_metadata(
    doc_hash: str,
) -> Tuple[Dict[str, Any], Optional[DocumentStructure], Optional[List[ChunkMetadata]]]:
    try:
        with open(Config.METADATA_CACHE_DIR / f"{doc_hash}.json", "r") as f:
            cache_data = json.load(f)

        metadata = cache_data.get("metadata", {})

        doc_structure = None
        if cache_data.get("doc_structure"):
            doc_structure = DocumentStructure()
            doc_structure.__dict__.update(cache_data["doc_structure"])

        chunk_metadata = None
        if cache_data.get("chunk_metadata"):
            chunk_metadata = []
            for cm_data in cache_data["chunk_metadata"]:
                cm = ChunkMetadata(
                    chunk_id=cm_data["chunk_id"],
                    page_num=cm_data.get("page_num"),
                    section=cm_data.get("section"),
                    keywords=cm_data.get("keywords", []),
                    chunk_type=cm_data.get("chunk_type", "text"),
                    importance_score=cm_data.get("importance_score", 1.0),
                )
                chunk_metadata.append(cm)

        return metadata, doc_structure, chunk_metadata
    except (FileNotFoundError, json.JSONDecodeError):
        return {}, None, None


# vector store creation
def create_vector_store(
    chunks: List[str],
    embeddings: SentenceTransformerEmbeddings,
    chunk_metadata: Optional[List[ChunkMetadata]] = None,
) -> FAISS:
    if not chunks:
        raise ValueError("Cannot create vector store from empty text chunks")

    logger.info(f"Creating vector store with {len(chunks)} chunks")
    try:
        # metadata for FAISS
        metadatas = []
        for i, _ in enumerate(chunks):
            metadata = {"chunk_index": i}
            if chunk_metadata and i < len(chunk_metadata):
                cm = chunk_metadata[i]
                metadata.update(
                    {
                        "chunk_id": cm.chunk_id,
                        "page_num": cm.page_num,
                        "section": cm.section,
                        "keywords": cm.keywords,
                        "chunk_type": cm.chunk_type,
                        "importance_score": cm.importance_score,
                    }
                )
            metadatas.append(metadata)

        vector_store = FAISS.from_texts(
            texts=chunks, embedding=embeddings, metadatas=metadatas
        )
        return vector_store
    except Exception as e:
        logger.error(f"Failed to create vector store: {str(e)}")
        raise DocumentProcessingError("Failed to generate document embeddings")


# context retrieval
def get_context(
    question: str,
    vector_store: FAISS,
    top_k: int,
    chunk_metadata: Optional[List[ChunkMetadata]] = None,
) -> Tuple[str, Dict[str, Any]]:
    # vector search
    primary_docs = vector_store.similarity_search(question, k=top_k)
    context_metadata = {
        "primary_chunks": len(primary_docs),
        "strategies_used": ["vector_similarity"],
        "total_chunks": len(primary_docs),
    }
    selected_chunks = [
        (doc.page_content, doc.metadata.get("importance_score", 1.0))
        for doc in primary_docs
    ]

    # keyword-based boost
    if chunk_metadata:
        question_keywords = extract_keywords(question.lower())

        # finding chunks with matching keywords
        keyword_matches = []
        all_texts = [
            doc.page_content
            for doc in vector_store.similarity_search(
                "", k=len(chunk_metadata) if chunk_metadata else 100
            )
        ]

        for i, text in enumerate(all_texts):
            if i < len(chunk_metadata):
                cm = chunk_metadata[i]
                # check keyword overlap
                text_keywords = [kw.lower() for kw in cm.keywords]
                question_kw_lower = [kw.lower() for kw in question_keywords]

                overlap = set(text_keywords).intersection(set(question_kw_lower))
                if overlap and text not in [chunk[0] for chunk in selected_chunks]:
                    boost_score = (
                        len(overlap) * Config.KEYWORD_BOOST_FACTOR * cm.importance_score
                    )
                    keyword_matches.append((text, boost_score, cm))

        keyword_matches.sort(key=lambda x: x[1], reverse=True)
        for text, score, cm in keyword_matches[:2]:  # top 2 matches
            selected_chunks.append((text, score))

        context_metadata["strategies_used"].append("keyword_matching")
        context_metadata["keyword_matches"] = len(keyword_matches[:2])

        # Context expansion - add neighboring chunks
        expanded_chunks = []
        primary_indices = [doc.metadata.get("chunk_index", -1) for doc in primary_docs]

        for idx in primary_indices:
            if idx > 0:  # previous chunk
                try:
                    prev_doc = vector_store.similarity_search("", k=idx + 1)[idx - 1]
                    if prev_doc.page_content not in [
                        chunk[0] for chunk in selected_chunks
                    ]:
                        expanded_chunks.append((prev_doc.page_content, 0.8))
                except:
                    pass

            # todo next chunk

        selected_chunks.extend(expanded_chunks[:2])  # upto 2 chunks
        context_metadata["strategies_used"].append("context_expansion")
        context_metadata["expanded_chunks"] = len(expanded_chunks[:2])

    # Sort by importance score and combine
    selected_chunks.sort(key=lambda x: x[1], reverse=True)
    context_metadata["total_chunks"] = len(selected_chunks)

    # Combine all selected chunks
    context_parts = []
    for i, (chunk_text, score) in enumerate(selected_chunks):
        context_parts.append(
            f"--- Context Chunk {i+1} (Score: {score:.2f}) ---\n{chunk_text}"
        )

    final_context = "\n\n".join(context_parts)

    return final_context, context_metadata


# LLM
async def generate_answer_with_llm(
    question: str, context: str, context_metadata: Dict[str, Any]
) -> str:
    strategies_used = ", ".join(context_metadata.get("strategies_used", []))
    total_chunks = context_metadata.get("total_chunks", 0)

    prompt = f"""
Role: You are an expert insurance claims adjudicator with deep knowledge of policy interpretation. Provide clear, accurate information based strictly on the policy document.

Context Analysis: 
    - Retrieved {total_chunks} relevant document sections using: {strategies_used}
    - Each section is scored by relevance and importance

Instructions:
    Communication Style:
        - Write in a single, clear paragraph that a customer can easily understand
        - Use exact policy wording when referencing specific terms, amounts, or time periods
        - Maintain original formatting of numbers and periods (e.g., "thirty (30) days")
        - Be conversational but precise - avoid bullet points or structured lists

    Content Requirements:
        - Start with a clear answer: "Yes, this is covered" / "No, this is not covered" / "This is partially covered" / "Based on the available information..."
        - Include the specific policy section reference (e.g., "According to Section 2.21...")
        - Mention any coverage amounts, limits, or time periods
        - Explain any conditions that must be met
        - Note any exclusions or limitations that apply
        - If information is incomplete, clearly state what additional details would be needed

    Quality Standards:
        - Never extrapolate beyond the provided policy text
        - Use policy-defined terms consistently
        - If there are contradictions or ambiguities, acknowledge them
        - Be honest about the completeness of the information available

Policy Context (Analyzed from {total_chunks} relevant sections):
{context}

Question: {question}

Required Output: Write a single, comprehensive paragraph that directly answers the customer's question with all relevant policy details, conditions, and limitations in a natural, flowing narrative format.
    """

    logger.info(f"Generating answer for question: {question[:50]}...")
    return await generate_text_with_retry(
        prompt=prompt,
        model_name=Config.GEMINI_MODEL,
        #max_retries=Config.MAX_LLM_RETRIES,
    )


# process questions in parallel
async def process_questions(
    questions: List[str],
    vector_store: FAISS,
    top_k: int,
    chunk_metadata: Optional[List[ChunkMetadata]] = None,
) -> Tuple[List[str], List[Dict[str, Any]]]:
    loop = asyncio.get_event_loop()
    tasks = []
    context_metadatas = []

    with ThreadPoolExecutor() as executor:
        # get context for all questions
        for question in questions:
            context, ctx_metadata = await loop.run_in_executor(
                executor,
                partial(
                    get_context,
                    question,
                    vector_store,
                    top_k,
                    chunk_metadata,
                ),
            )
            context_metadatas.append(ctx_metadata)
            tasks.append(generate_answer_with_llm(question, context, ctx_metadata))

        answers = await asyncio.gather(*tasks)
        return answers, context_metadatas


# API endpoints
@app.get("/")
async def root():
    return {"message": "Letsgoo, its up running........"}


@app.post("/api/v1/hackrx/run", response_model=HackRxResponse)
async def hackrx_run(
    req: HackRxRequest, api_key: str = Security(get_api_key)
) -> HackRxResponse:
    start_time = asyncio.get_event_loop().time()

    try:
        content_bytes, file_ext = await get_document_content_from_url(req.documents)
        doc_hash = get_document_hash(content_bytes.getvalue())
        vector_store_path = Config.VECTORDB_CACHE_DIR / f"{doc_hash}.faiss"

        # cache check
        if vector_store_path.exists():
            logger.info(f"Loading vector store from cache: {vector_store_path}")
            vector_store = FAISS.load_local(
                folder_path=str(Config.VECTORDB_CACHE_DIR),
                index_name=doc_hash,
                embeddings=app.state.embeddings,
                allow_dangerous_deserialization=True,
            )

            # load metadata
            metadata, doc_structure, chunk_metadata = load_metadata(doc_hash)
            if not metadata:
                metadata = {"source": "cache"}
        else:
            logger.info("Processing new document.")

            parser = PARSER_MAPPING.get(file_ext)
            if not parser:
                raise HTTPException(
                    status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                    detail=f"File type '{file_ext}' not supported",
                )

            (
                raw_text,
                metadata,
                doc_structure,
            ) = await asyncio.get_event_loop().run_in_executor(
                None, parser, io.BytesIO(content_bytes.getvalue())
            )

            text_chunks, chunk_metadata = clean_and_chunk_text(
                raw_text, req.chunk_size, req.chunk_overlap
            )

            vector_store = await asyncio.get_event_loop().run_in_executor(
                None,
                create_vector_store,
                text_chunks,
                app.state.embeddings,
                chunk_metadata,
            )

            # save to cache
            vector_store.save_local(
                folder_path=str(Config.VECTORDB_CACHE_DIR), index_name=doc_hash
            )
            save_metadata(doc_hash, metadata, doc_structure, chunk_metadata)

            with open(Config.DOCUMENT_CACHE_DIR / doc_hash, "wb") as f:
                f.write(content_bytes.getvalue())
            logger.info(
                f"Saved new document and vector store to cache with hash: {doc_hash}"
            )

        if chunk_metadata:
            # process questions with context retrieval
            answers, context_metadatas = await process_questions(
                req.questions,
                vector_store,
                req.top_k,
                chunk_metadata,
            )

            # mesh it up
            aggregated_context_metadata = {
                "total_questions": len(req.questions),
                "average_chunks_per_question": sum(
                    cm.get("total_chunks", 0) for cm in context_metadatas
                )
                / len(context_metadatas),
                "strategies_summary": {
                    strategy: sum(
                        1
                        for cm in context_metadatas
                        if strategy in cm.get("strategies_used", [])
                    )
                    for strategy in [
                        "vector_similarity",
                        "keyword_matching",
                        "context_expansion",
                    ]
                },
            }
        else:
            answers = await process_questions(req.questions, vector_store, req.top_k, None)
            aggregated_context_metadata = {
                "total_questions": len(req.questions),
                "processing_mode": "standard",
            }

        processing_time = asyncio.get_event_loop().time() - start_time

        # metadata for response
        response_metadata = metadata.copy()
        if doc_structure:
            response_metadata.update(
                {
                    "document_structure": {
                        "sections_found": len(doc_structure.sections),
                        "headers_found": len(doc_structure.headers),
                        "tables_found": len(doc_structure.tables),
                        "key_value_pairs": len(doc_structure.key_value_pairs),
                    }
                }
            )

        if chunk_metadata:
            chunk_types = {}
            for cm in chunk_metadata:
                chunk_types[cm.chunk_type] = chunk_types.get(cm.chunk_type, 0) + 1
            response_metadata["chunk_analysis"] = {
                "total_chunks": len(chunk_metadata),
                "chunk_types": chunk_types,
                "avg_importance_score": sum(
                    cm.importance_score for cm in chunk_metadata
                )
                / len(chunk_metadata),
            }

        return HackRxResponse(
            answers=answers,
            processing_time=round(processing_time, 2),
            document_metadata=response_metadata,
            context_metadata=aggregated_context_metadata,
        )

    except (HTTPException, DocumentProcessingError, LLMGenerationError) as e:
        raise e
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred",
        )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8080,
        proxy_headers=True,
        forwarded_allow_ips="*",
        log_config=None,
        reload=True if os.getenv("ENVIRONMENT") == "development" else False,
    )
