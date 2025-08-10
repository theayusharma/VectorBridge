import asyncio
import concurrent.futures
import io
import logging
import os
import re
import tempfile
from abc import abstractmethod
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from threading import Lock
from typing import Any, Dict, List, Optional, Set, Tuple

import camelot
import docx
import pdfplumber
import textract

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    pass


@dataclass
class DocumentStructure:
    headers: List[Tuple[str, int]]
    sections: Dict[str, List[str]]
    tables: List[Dict[str, Any]]
    key_value_pairs: Dict[str, str]

    def __init__(self):
        self.headers = []
        self.sections = {}
        self.tables = []
        self.key_value_pairs = {}


class BaseParser:
    @abstractmethod
    def parse(
        self, content: io.BytesIO
    ) -> Tuple[str, Dict[str, Any], Optional[DocumentStructure]]:
        pass


class PDFParser(BaseParser):
    def __init__(self, max_table_workers: int = 4):
        super().__init__()
        self.max_table_workers = max_table_workers
        self._camelot_lock = Lock()

    def parse(self, content: io.BytesIO) -> Tuple[str, Dict[str, Any], DocumentStructure]:
        working_copy = io.BytesIO(content.getvalue())
        full_text = []
        metadata = {"pages": 0, "tables": 0, "sections": 0}
        doc_structure = DocumentStructure()

        try:
            # First pass - identify pages with potential tables
            with pdfplumber.open(working_copy) as pdf:
                metadata["pages"] = len(pdf.pages)
                table_pages = self._identify_table_pages(pdf)
                
                # Process text content sequentially
                current_section = "Introduction"
                section_text = []
                
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
                    self._process_page_content(
                        page_text, i, doc_structure, metadata, 
                        full_text, current_section, section_text
                    )

            # Parallel table extraction
            if table_pages:
                working_copy.seek(0)
                tables = self._parallel_table_extraction(working_copy, table_pages)
                metadata["tables"] = len(tables)
                self._process_tables(tables, full_text, doc_structure)

            return "\n".join(full_text), metadata, doc_structure

        except Exception as e:
            raise DocumentProcessingError(f"PDF processing failed: {str(e)}")
        finally:
            working_copy.close()

    def _identify_table_pages(self, pdf) -> Set[int]:
        """Identify pages likely containing tables"""
        table_pages = set()
        for i, page in enumerate(pdf.pages):
            if self._page_likely_has_tables(page):
                table_pages.add(i + 1)  # Camelot uses 1-based page numbers
        return table_pages

    def _is_header(self, text: str) -> bool:
        """Improved header detection logic"""
        if not text:
            return False
        # Check for all caps with reasonable length
        if text.isupper() and 2 < len(text) < 80:
            return True
        # Check for numbered headings (e.g., "1. Introduction")
        if re.match(r"^\d+\.\s+\w+", text):
            return True
        # Check for section headings ending with colon
        if text.endswith(":") and len(text.split()) < 8:
            return True
        return False

    def _determine_header_level(self, text: str) -> int:
        """Determine heading level based on formatting"""
        if text.isupper():
            return 1  # Highest level
        if re.match(r"^\d+\.", text):
            return 2
        return 3  # Default level
    def _process_page_content(self, page_text: str, page_num: int, doc_structure: DocumentStructure,
                             metadata: Dict[str, Any], full_text: List[str],
                             current_section: str, section_text: List[str]) -> str:
        """Process the content of a single page and update document structure"""
        lines = page_text.split("\n")
        
        for line in lines:
            line = line.strip()
            if self._is_header(line):
                header_level = self._determine_header_level(line)
                doc_structure.headers.append((line, header_level))
                current_section = line
                metadata["sections"] += 1
                # Flush previous section content
                if section_text:
                    doc_structure.sections[current_section] = section_text
                    section_text = []
    
        section_text.append(page_text)
        full_text.append(
            f"--- Page {page_num+1} | Section: {current_section} ---\n{page_text}"
        )
        
        # Extract key-value pairs
        self._extract_key_value_pairs(page_text, doc_structure)
        
        return current_section

    def _extract_key_value_pairs(self, text: str, doc_structure: DocumentStructure):
        """Extract key: value pairs from text"""
        kv_pattern = re.compile(r"([A-Z][A-Za-z\s]+):\s*([^\n]+)")
        for key, value in kv_pattern.findall(text):
            clean_key = key.strip()
            clean_value = value.strip()
            if clean_key and clean_value:
                doc_structure.key_value_pairs[clean_key] = clean_value

    def _parallel_table_extraction(self, pdf_stream: io.BytesIO, pages: Set[int]) -> List[Any]:
        """Extract tables from multiple pages in parallel"""
        tables = []
        page_numbers = list(pages)
        
        # We process pages in chunks to balance memory usage
        chunk_size = max(1, len(page_numbers) // self.max_table_workers)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_table_workers) as executor:
            futures = []
            
            for i in range(0, len(page_numbers), chunk_size):
                chunk = page_numbers[i:i + chunk_size]
                futures.append(
                    executor.submit(
                        self._extract_tables_chunk,
                        pdf_stream.getvalue(),  # Pass raw bytes to avoid thread issues
                        chunk
                    )
                )
            
            for future in concurrent.futures.as_completed(futures):
                try:
                    tables.extend(future.result())
                except Exception as e:
                    logger.error(f"Table extraction failed for chunk: {str(e)}")
        
        return tables

    def _extract_tables_chunk(self, pdf_bytes: bytes, pages: List[int]) -> List[Any]:
        """Process a chunk of pages (runs in worker thread)"""
        chunk_tables = []
        pdf_stream = io.BytesIO(pdf_bytes)
        
        try:
            # Camelot isn't thread-safe, so we need to lock
            with self._camelot_lock:
                tables = camelot.read_pdf(
                    pdf_stream,
                    flavor="lattice",
                    pages=",".join(map(str, pages)),
                    suppress_stdout=True,
                    process_background=True
                )
                chunk_tables.extend(tables)
        finally:
            pdf_stream.close()
        
        return chunk_tables

    def _process_tables(self, tables, full_text, doc_structure):
        """Process extracted tables into structured format"""
        for table in tables:
            if table.shape[0] <= 1:
                continue

            table_metadata = {
                "page": table.page,
                "table_id": f"table_{table.page}_{table.order}",
                "rows": table.shape[0],
                "cols": table.shape[1],
            }
            doc_structure.tables.append(table_metadata)

            # Convert to markdown
            df = table.df
            headers = [self._clean_table_cell(cell) for cell in df.iloc[0].tolist()]
            rows = [
                "| " + " | ".join(self._clean_table_cell(cell) for cell in row) + " |"
                for _, row in df[1:].iterrows()
            ]

            markdown_table = (
                "| " + " | ".join(headers) + " |\n"
                "| "
                + " | ".join(["---"] * len(headers))
                + " |\n"
                + "\n".join(rows)
                + "\n"
            )

            full_text.append(
                f"\n--- Table {table.order} on Page {table.page} ---\n{markdown_table}\n"
            )

    def _clean_table_cell(self, text: str) -> str:
        """Clean table cell content"""
        if not text or not isinstance(text, str):
            return ""
        return re.sub(r"\s+", " ", str(text).strip())


    def _page_likely_has_tables(self, page: pdfplumber.page) -> bool:
        """Heuristic to determine if a page likely contains tables"""
        # Check for common table indicators
        text = page.extract_text() or ""
        
        # 1. Look for dense text areas with alignment patterns
        words = page.extract_words()
        if len(words) > 50:  # Dense text
            x_coords = [w['x0'] for w in words]
            # Check for alignment patterns (common in tables)
            if len(set(round(x) for x in x_coords)) < len(x_coords)/3:
                return True
        
        # 2. Look for common table headers/footers
        table_keywords = ["table", "tab.", "figure", "no.", "item", "description", "qty", "amount"]
        if any(keyword in text.lower() for keyword in table_keywords):
            return True
        
        # 3. Look for grid-like structures
        if len(page.rect_edges) > 10:  # Lots of rectangles
            return True
        
        return False


class DOCXParser(BaseParser):
    def parse(
        self, content: io.BytesIO
    ) -> Tuple[str, Dict[str, Any], DocumentStructure]:
        try:
            doc = docx.Document(content)
            full_text = []
            current_section = "Introduction"
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
            }
            doc_structure = DocumentStructure()

            for para in doc.paragraphs:
                if para.text.strip():  # todo check errors from lsp
                    # header
                    if para.style.name.startswith("Heading") or (
                        para.text.isupper() and len(para.text) < 100
                    ):
                        level = 1
                        if "Heading" in para.style.name:
                            level = (
                                int(para.style.name.split()[-1])
                                if para.style.name.split()[-1].isdigit()
                                else 1
                            )
                        doc_structure.headers.append((para.text, level))
                        current_section = para.text

                    # add to section
                    if current_section not in doc_structure.sections:
                        doc_structure.sections[current_section] = []
                    doc_structure.sections[current_section].append(para.text)

                    full_text.append(f"[Section: {current_section}] {para.text}")

            # tables
            for i, table in enumerate(doc.tables):
                table_text = [
                    " | ".join(cell.text.strip() for cell in row.cells)
                    for row in table.rows
                ]

                table_metadata = {
                    "table_id": f"docx_table_{i+1}",
                    "rows": len(table.rows),
                    "cols": len(table.rows[0].cells) if table.rows else 0,
                }
                doc_structure.tables.append(table_metadata)
                full_text.append(
                    f"\n--- Table {i+1} ---\n" + "\n".join(table_text) + "\n"
                )

            return "\n".join(full_text), metadata, doc_structure
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse DOCX: {str(e)}") from e


class DOCParser(BaseParser):
    def parse(self, content: io.BytesIO) -> Tuple[str, Dict[str, Any], None]:
        try:
            with tempfile.NamedTemporaryFile(delete=True) as tmp:
                tmp.write(content.read())
                tmp.flush()
                text = textract.process(tmp.name).decode("utf-8")
            return text, {"format": "DOC"}, None
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse DOC: {str(e)}") from e


class TXTParser(BaseParser):
    def parse(self, content: io.BytesIO) -> Tuple[str, Dict[str, Any], None]:
        try:
            text = content.read().decode("utf-8", errors="ignore")
            return text, {"size": len(text)}, None
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse TXT: {str(e)}") from e


# parser mapping
PARSER_MAPPING: Dict[str, BaseParser] = {
    ".pdf": PDFParser(),
    ".docx": DOCXParser(),
    ".doc": DOCParser(),
    ".txt": TXTParser(),
}


class DocumentParser:
    def __init__(
        self,
        parsers: Dict[str, Any],
        max_workers: int = 4,
        max_file_size: int = 30 * 1024 * 1024,  # 30MB
        timeout: float = 30.0,
    ):
        self.parsers = parsers
        self.executor = ThreadPoolExecutor(max_workers=max_workers)
        self.max_file_size = max_file_size
        self.timeout = timeout

    async def parse(
        self, content: io.BytesIO, file_ext: str
    ) -> Tuple[str, Dict[str, Any], Optional[DocumentStructure]]:
        logger.info("Parsing a doc")
        # Validate input
        content.seek(0, os.SEEK_END)
        file_size = content.tell()
        content.seek(0)

        if file_size > self.max_file_size:
            raise DocumentProcessingError(
                f"File size {file_size/1024/1024:.2f}MB exceeds maximum {self.max_file_size/1024/1024:.2f}MB"
            )

        # get appropriate parser
        parser = self._get_parser(file_ext)

        try:
            # run parsing in thread pool with timeout
            loop = asyncio.get_event_loop()
            result = await asyncio.wait_for(
                loop.run_in_executor(
                    self.executor, self._parse_with_retry, parser, content
                ),
                timeout=self.timeout,
            )
            return result
        except asyncio.TimeoutError:
            raise DocumentProcessingError(
                f"Parsing timed out after {self.timeout} seconds"
            )
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse document: {str(e)}")

    def _get_parser(self, file_ext: str) -> BaseParser:
        """Get parser instance for file extension"""
        file_ext = file_ext.lower()
        if file_ext not in self.parsers:
            if file_ext in [".rtf", ".odt"]:  # fallbacks
                return self.parsers.get(".docx", self.parsers[".txt"])
            return self.parsers[".txt"]  # default fallback
        return self.parsers[file_ext]

    def _parse_with_retry(
        self, parser: BaseParser, content: io.BytesIO
    ) -> Tuple[str, Dict[str, Any], Optional[DocumentStructure]]:
        """Parse document with retries"""
        max_attempts = 3
        last_exception = None

        for attempt in range(1, max_attempts + 1):
            try:
                content_copy = io.BytesIO(content.getvalue())
                return parser.parse(content_copy)
            except Exception as e:
                last_exception = e
                logger.warning(f"Parse attempt {attempt} failed: {str(e)}")
                if attempt < max_attempts:
                    continue
                raise DocumentProcessingError(
                    f"Failed after {max_attempts} attempts: {str(last_exception)}"
                )

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.executor.shutdown(wait=True)

    async def close(self):
        """Cleanup resources"""
        self.executor.shutdown(wait=True)
